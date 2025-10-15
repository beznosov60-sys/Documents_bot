from __future__ import annotations

import logging
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from jinja2 import Template
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from bot.models import ContractContext, Payment
from bot.utils.formatting import amount_to_words, format_russian_date

logger = logging.getLogger(__name__)

DEFAULT_SECTIONS = {
    "section_1": "Исполнитель обязуется оказать юридические услуги в интересах Заказчика в соответствии с условиями настоящего договора.",
    "section_2": "Заказчик предоставляет Исполнителю всю необходимую информацию и документы для оказания услуг, а Исполнитель обязуется сохранять конфиденциальность.",
    "section_3": "Оплата услуг осуществляется в соответствии с графиком платежей, являющимся неотъемлемой частью настоящего договора.",
    "section_4": "Настоящий договор вступает в силу с момента его подписания и действует до полного исполнения сторонами обязательств.",
}

SECTION_TITLES = [
    "Общие положения",
    "Права и обязанности сторон",
    "Порядок расчетов",
    "Заключительные положения",
]

PDF_FONT_NAME = "TimesNewRoman"

INTRO_TEMPLATE = Template(
    "<b>Исполнитель:</b> Заруцкая К.Н., действующая на основании Устава, с одной стороны, и Заказчик: "
    "<b>{{ passport.full_name }}</b>, с паспортом серии {{ passport.series }} номер {{ passport.number }}, "
    "выданным {{ passport.issued_by }} {{ issued_date }}, заключили настоящий договор о нижеследующем."
)


async def generate_documents(context: ContractContext, destination_dir: Path) -> ContractContext:
    destination_dir.mkdir(parents=True, exist_ok=True)
    base_filename = f"dogovor_{context.contract_number}_{context.passport.full_name.replace(' ', '_')}"
    docx_path = destination_dir / f"{base_filename}.docx"
    pdf_path = destination_dir / f"{base_filename}.pdf"

    _create_docx(context, docx_path)
    _create_pdf(pdf_path, context)

    context.contract_dir = destination_dir
    context.docx_path = docx_path
    context.pdf_path = pdf_path
    return context


def _create_docx(context: ContractContext, docx_path: Path) -> None:
    document = Document()
    if document.paragraphs:
        first = document.paragraphs[0]._element
        first.getparent().remove(first)

    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    initials = "".join(part[0] for part in context.passport.full_name.split() if part)

    _add_paragraph(
        document,
        "ДОГОВОР",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    _add_paragraph(
        document,
        f"об оказании юридических услуг №{context.contract_number} БФЛ 127 ФЗ {initials}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )

    _add_paragraph(document, "г. Москва", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, format_russian_date(date.today()), alignment=WD_ALIGN_PARAGRAPH.CENTER)

    intro_text = INTRO_TEMPLATE.render(
        passport=context.passport,
        issued_date=format_russian_date(context.passport.issued_date),
    )
    _add_html_paragraph(document, intro_text)

    _add_paragraph(
        document,
        "Общая сумма договора составляет "
        f"{context.total_amount:,.0f} ({amount_to_words(context.total_amount)}) рублей.",
    )

    for index, title in enumerate(SECTION_TITLES, start=1):
        _add_paragraph(
            document,
            f"{index}. {title}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        _add_paragraph(document, DEFAULT_SECTIONS[f"section_{index}"])

    _add_paragraph(document, "График платежей:", bold=True)
    _insert_payment_table(document, context.payments)

    document.save(docx_path)


def _add_paragraph(
    document: Document,
    text: str,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    bold: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.15


def _add_html_paragraph(document: Document, html_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    remaining = html_text
    while remaining:
        bold_start = remaining.find("<b>")
        if bold_start == -1:
            run = paragraph.add_run(remaining)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            break
        if bold_start > 0:
            run = paragraph.add_run(remaining[:bold_start])
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        bold_end = remaining.find("</b>", bold_start)
        if bold_end == -1:
            run = paragraph.add_run(remaining[bold_start:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            break
        bold_text = remaining[bold_start + 3 : bold_end]
        run = paragraph.add_run(bold_text)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        remaining = remaining[bold_end + 4 :]


def _insert_payment_table(document: Document, payments: Iterable[Payment]) -> None:
    payments = list(payments)
    table = document.add_table(rows=1 + len(payments), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Месяц", "Дата платежа", "Сумма, ₽"]
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    for row_index, payment in enumerate(payments, start=1):
        row = table.rows[row_index]
        values = [
            str(payment.month_index),
            format_russian_date(payment.due_date),
            f"{payment.amount:,.0f}".replace(",", " "),
        ]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def _register_pdf_font() -> str:
    possible_fonts = [
        Path("/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf"),
        Path("/usr/share/fonts/truetype/msttcorefonts/times.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"),
        Path("/usr/share/fonts/truetype/freefont/FreeSerif.ttf"),
    ]
    for font_path in possible_fonts:
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(font_path)))
                logger.debug("Registered PDF font from %s", font_path)
                return PDF_FONT_NAME
            except Exception as exc:  # pragma: no cover - font registration edge case
                logger.warning("Failed to register font %s: %s", font_path, exc)
    return "Times-Roman"


def _create_pdf(pdf_path: Path, context: ContractContext) -> None:
    font_name = _register_pdf_font()
    logger.debug("Generating PDF at %s using font %s", pdf_path, font_name)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
    )

    styles = {
        "title": ParagraphStyle(
            name="Title",
            fontName=font_name,
            fontSize=14,
            leading=16,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            name="Subtitle",
            fontName=font_name,
            fontSize=12,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "normal": ParagraphStyle(
            name="Normal",
            fontName=font_name,
            fontSize=12,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
        ),
    }

    passport = context.passport
    story = []

    title_text = (
        f"ДОГОВОР<br/>об оказании юридических услуг №{context.contract_number} БФЛ 127 ФЗ {''.join(part[0] for part in passport.full_name.split() if part)}"
    )
    story.append(Paragraph(title_text, styles["title"]))
    story.append(Paragraph("г. Москва", styles["subtitle"]))
    story.append(Paragraph(format_russian_date(date.today()), styles["subtitle"]))

    intro_text = (
        "<b>Исполнитель:</b> Заруцкая К.Н., действующая на основании Устава, с одной стороны, и Заказчик: "
        f"<b>{passport.full_name}</b>, с паспортом серии {passport.series} номер {passport.number}, "
        f"выданным {passport.issued_by} {format_russian_date(passport.issued_date)}, заключили настоящий договор о нижеследующем."
    )
    story.append(Paragraph(intro_text, styles["normal"]))

    summary_text = (
        "Общая сумма договора составляет "
        f"{context.total_amount:,.0f} ({amount_to_words(context.total_amount)}) рублей."
    )
    story.append(Paragraph(summary_text, styles["normal"]))

    for index, title in enumerate(SECTION_TITLES, start=1):
        header = ParagraphStyle(
            name=f"Header{index}",
            parent=styles["normal"],
            fontSize=12,
            alignment=TA_CENTER,
            leading=14,
        )
        story.append(Paragraph(f"<b>{index}. {title}</b>", header))
        story.append(Paragraph(DEFAULT_SECTIONS[f"section_{index}"] , styles["normal"]))

    story.append(Spacer(1, 12))

    table_data = [["Месяц", "Дата платежа", "Сумма, ₽"]]
    for payment in context.payments:
        table_data.append([
            str(payment.month_index),
            format_russian_date(payment.due_date),
            f"{payment.amount:,.0f}".replace(",", " "),
        ])

    payment_table = Table(table_data, colWidths=[3 * cm, 6 * cm, 4 * cm])
    payment_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONT", (0, 0), (-1, -1), font_name, 12),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]
        )
    )
    story.append(payment_table)

    doc.build(story)
